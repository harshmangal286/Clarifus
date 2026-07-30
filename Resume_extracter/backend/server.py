"""Clarifuse's local resume-extraction API.

Run from the project root with: python -m uvicorn backend.server:app --reload
"""
from __future__ import annotations

import io
import json
import os
import re
from pathlib import Path
from typing import List, Optional

import requests
from docx import Document
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from PyPDF2 import PdfReader

ROOT = Path(__file__).parent.parent
FRONTEND = ROOT / "frontend" / "dist"

app = FastAPI(title="Clarifuse Resume Extraction")

# Try to mount frontend assets if they exist (used for production serving)
if (FRONTEND / "assets").exists():
    app.mount("/assets", StaticFiles(directory=FRONTEND / "assets"), name="assets")

# ==========================================
# 7-LAYER RESUME EXTRACTION SCHEMAS
# ==========================================

class EducationInfo(BaseModel):
    current_year: Optional[str] = Field(None, description="Current year of study, e.g., '3rd Year', 'Final Year', or null")
    college: Optional[str] = Field(None, description="Name of the college/university")
    degree: Optional[str] = Field(None, description="Degree name, e.g., 'B.Tech in Computer Science'")
    cgpa: Optional[float] = Field(None, description="Current CGPA or percentage")
    graduation_year: Optional[int] = Field(None, description="Expected graduation year, e.g., 2025")

class PersonalInfo(BaseModel):
    name: str = Field(..., description="Full name of the student")
    email: Optional[str] = Field(None, description="Email address")
    phone: Optional[str] = Field(None, description="Contact phone number")
    linkedin: Optional[str] = Field(None, description="LinkedIn profile URL")
    github: Optional[str] = Field(None, description="GitHub profile URL")
    portfolio: Optional[str] = Field(None, description="Portfolio website/blog URL")
    location: Optional[str] = Field(None, description="Location, e.g. 'City, State/Country'")
    current_education: Optional[EducationInfo] = Field(None, description="Details of their current education program")

class Skill(BaseModel):
    skill: str = Field(..., description="Name of the skill, e.g., React, Python")
    category: str = Field(..., description="Category, e.g., Frontend, Backend, AI/ML, DevOps, Database, Languages, Soft Skills")
    evidence: List[str] = Field(..., description="List of specific projects, internships, or experiences in this resume where this skill was applied/proven")
    confidence: float = Field(..., description="Confidence score between 0.0 and 1.0 based on depth of evidence")

class Project(BaseModel):
    title: str = Field(..., description="Title of the project")
    description: str = Field(..., description="Brief summary of the project and its goals")
    duration: Optional[str] = Field(None, description="Duration, e.g., 'Oct 2023 - Dec 2023'")
    team_size: Optional[str] = Field(None, description="Size of the project team, e.g., '1 (Solo)', '4 members'")
    role: Optional[str] = Field(None, description="Role in the project, e.g., 'Frontend Lead', 'Fullstack Developer'")
    tech_stack: List[str] = Field(default_factory=list, description="Technologies and libraries used")
    features: List[str] = Field(default_factory=list, description="Key features built in this project")
    domain: List[str] = Field(default_factory=list, description="Domains this project belongs to, e.g., Backend, Frontend, AI, ML, IoT, Cloud, Blockchain, Cybersecurity, Data Science")
    deployment: Optional[str] = Field(None, description="Where the project is deployed, e.g., AWS, Vercel, Netlify, Github Pages, or null")
    github: Optional[str] = Field(None, description="GitHub repository link for this project")
    live_demo: Optional[str] = Field(None, description="Live demo or website link")
    impact: Optional[str] = Field(None, description="Real-world impact, metrics, or user value achieved")
    complexity: Optional[str] = Field(None, description="Inferred complexity: High, Medium, or Low, with a brief explanation")
    keywords: List[str] = Field(default_factory=list, description="Keywords for indexing and filtering")

class WorkExperience(BaseModel):
    company: str = Field(..., description="Company or organization name")
    position: str = Field(..., description="Job title / position")
    duration: str = Field(..., description="Duration of employment, e.g., 'May 2023 - July 2023'")
    employment_type: Optional[str] = Field(None, description="Employment type, e.g. 'Internship', 'Full-time', 'Freelance'")
    responsibilities: List[str] = Field(default_factory=list, description="Key duties and tasks performed")
    technologies: List[str] = Field(default_factory=list, description="Technologies used during this experience")
    achievements: List[str] = Field(default_factory=list, description="Key accomplishments or successful implementations")
    leadership: List[str] = Field(default_factory=list, description="Leadership roles, mentoring, or initiatives taken")
    impact: List[str] = Field(default_factory=list, description="Measurable impact, improvements, or outcomes")

class Education(BaseModel):
    college: str = Field(..., description="Name of the institution (school or college)")
    degree: Optional[str] = Field(None, description="Degree obtained or pursuing")
    department: Optional[str] = Field(None, description="Department or field of study, e.g., Computer Science")
    year: Optional[str] = Field(None, description="Time duration, e.g., '2021 - 2025'")
    cgpa: Optional[str] = Field(None, description="CGPA or percentage grade obtained")
    relevant_coursework: List[str] = Field(default_factory=list, description="List of relevant courses taken")
    academic_honors: List[str] = Field(default_factory=list, description="Honors, Dean's list, scholarships, etc.")
    expected_graduation: Optional[str] = Field(None, description="Expected graduation date, e.g., 'May 2025'")

class Certification(BaseModel):
    name: str = Field(..., description="Name of the certification")
    issuer: str = Field(..., description="Issuing organization, e.g., AWS, Google, Udemy")
    issue_date: Optional[str] = Field(None, description="Date of issue")
    credential_url: Optional[str] = Field(None, description="URL to verify credentials")
    skills: List[str] = Field(default_factory=list, description="Skills validated by this certification")

class Extras(BaseModel):
    hackathons: List[str] = Field(default_factory=list, description="Participated hackathons and details")
    open_source: List[str] = Field(default_factory=list, description="Open source contributions, pull requests, repositories")
    competitions: List[str] = Field(default_factory=list, description="Competitions and hackathons participated")
    research_papers: List[str] = Field(default_factory=list, description="Research papers written or under review")
    publications: List[str] = Field(default_factory=list, description="Articles, publications, blog posts")
    volunteer_work: List[str] = Field(default_factory=list, description="Volunteer experiences and roles")
    leadership: List[str] = Field(default_factory=list, description="Leadership roles, mentoring, club heads")
    positions_of_responsibility: List[str] = Field(default_factory=list, description="Positions of responsibility held")
    achievements: List[str] = Field(default_factory=list, description="Key personal/academic achievements")
    awards: List[str] = Field(default_factory=list, description="Awards, scholarships, accolades won")
    languages: List[str] = Field(default_factory=list, description="Languages spoken")
    interests: List[str] = Field(default_factory=list, description="Personal hobbies and interests")
    clubs: List[str] = Field(default_factory=list, description="Clubs participated in")
    societies: List[str] = Field(default_factory=list, description="Societies participated in")

class ResumeProfile(BaseModel):
    personal_information: PersonalInfo = Field(..., description="Layer 1: Personal Identification & Current Education Details")
    skills: List[Skill] = Field(default_factory=list, description="Layer 2: Rich skills with evidence and confidence")
    projects: List[Project] = Field(default_factory=list, description="Layer 3: Detailed projects with domains, complexity, and impact")
    work_experience: List[WorkExperience] = Field(default_factory=list, description="Layer 4: Professional work history and achievements")
    education: List[Education] = Field(default_factory=list, description="Layer 5: Academic history, coursework, and honors")
    certifications: List[Certification] = Field(default_factory=list, description="Layer 6: Validated certifications and URLs")
    extras: Extras = Field(default_factory=dict, description="Layer 7: Extra activities, hackathons, open source, and achievements")


# ==========================================
# ENV VARIABLES AND UTILITIES
# ==========================================

def load_dotenv():
    """Find and load .env variables locally without external library."""
    env_paths = [
        Path(__file__).parent / ".env",
        Path(__file__).parent.parent / ".env",
        Path.cwd() / ".env"
    ]
    for env_path in env_paths:
        if env_path.exists():
            try:
                loaded = []
                for line in env_path.read_text(encoding="utf-8").splitlines():
                    line = line.strip()

                    if not line or line.startswith("#"):
                        continue

                    if "=" not in line:
                        continue

                    key, value = line.split("=", 1)
                    key = key.strip()
                    value = value.strip().strip('"').strip("'")

                    if value:
                        os.environ[key] = value
                        obfuscated = value[:10] + "..." if len(value) > 10 else "***"
                        loaded.append(f"{key} ({obfuscated})")
                    else:
                        # Clean up keys that are set to empty
                        os.environ.pop(key, None)
                
                if loaded:
                    print(f"INFO: Loaded environment variables from {env_path.name}: {', '.join(loaded)}")
                break
            except Exception as e:
                print(f"WARNING: Failed to parse {env_path}: {e}")

# Load environmental variables
load_dotenv()


def resolve_refs(schema: dict, defs: dict = None) -> dict:
    """Resolve Pydantic OpenAPI JSON Schema definitions ($defs) to form a flat schema for Gemini API."""
    if defs is None:
        defs = schema.get("$defs", schema.get("definitions", {}))
    
    if isinstance(schema, dict):
        if "$ref" in schema:
            ref_path = schema["$ref"]
            ref_key = ref_path.split("/")[-1]
            ref_schema = defs.get(ref_key, {})
            return resolve_refs(ref_schema, defs)
        
        resolved = {}
        for k, v in schema.items():
            if k in ("$defs", "definitions"):
                continue
            resolved[k] = resolve_refs(v, defs)
        return resolved
    
    elif isinstance(schema, list):
        return [resolve_refs(item, defs) for item in schema]
    
    return schema


def prepare_schema_for_gemini(pydantic_model) -> dict:
    """Flatten and format the Pydantic schema for the Gemini API structured output configuration."""
    schema = pydantic_model.model_json_schema()
    resolved = resolve_refs(schema)
    
    def clean_types(item):
        if isinstance(item, dict):
            cleaned = {}
            for k, v in item.items():
                if k == "type" and isinstance(v, str):
                    cleaned[k] = v.upper()
                elif k in ("title", "default"):
                    continue
                else:
                    cleaned[k] = clean_types(v)
            return cleaned
        elif isinstance(item, list):
            return [clean_types(x) for x in item]
        return item
    
    return clean_types(resolved)


def parse_json_response(text: str) -> dict:
    """Clean markdown code block markers and safely load JSON data."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    return json.loads(text)


# ==========================================
# OPENROUTER INTEGRATION (PRIMARY)
# ==========================================

def call_openrouter_api(prompt_text: str) -> dict:
    """Call OpenRouter API to extract structured resume details with model fallbacks."""
    api_key = os.environ.get("OPENROUTER_API_KEY") or os.environ.get("OPENROUTER_APIKEY")
    if not api_key:
        raise ValueError("OPENROUTER_API_KEY is not configured.")
        
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:8000",
        "X-Title": "Clarifuse Resume Extracter"
    }
    
    # Primary model is gemini-2.5-flash, falling back to openai or llama models
    models = ["google/gemini-2.5-flash", "openai/gpt-4o-mini", "meta-llama/llama-3.3-70b-instruct"]
    last_error = None
    
    for model in models:
        payload = {
            "model": model,
            "messages": [
                {
                    "role": "user",
                    "content": prompt_text
                }
            ],
            "response_format": {
                "type": "json_object"
            },
            "max_tokens": 4000
        }
        try:
            response = requests.post(url, headers=headers, json=payload, timeout=60)
            if response.status_code == 200:
                response_data = response.json()
                text_content = response_data["choices"][0]["message"]["content"]
                return parse_json_response(text_content)
            else:
                last_error = f"Model {model} failed (Status {response.status_code}): {response.text}"
        except Exception as e:
            last_error = f"Model {model} exception: {str(e)}"
            
    raise Exception(f"All OpenRouter models failed. Details: {last_error}")


# ==========================================
# GEMINI DIRECT INTEGRATION (SECONDARY BACKUP)
# ==========================================

def call_gemini_api(prompt_text: str, schema: dict) -> dict:
    """Call Gemini direct REST API to extract structured resume details."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("GEMINI_API_KEY is not configured.")
    
    models = ["gemini-2.5-flash", "gemini-1.5-flash", "gemini-2.5-pro"]
    last_error = None
    
    for model in models:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
        headers = {"Content-Type": "application/json"}
        payload = {
            "contents": [
                {
                    "parts": [
                        {
                            "text": prompt_text
                        }
                    ]
                }
            ],
            "generationConfig": {
                "responseMimeType": "application/json",
                "responseSchema": schema
            }
        }
        try:
            response = requests.post(url, headers=headers, json=payload, timeout=60)
            if response.status_code == 200:
                response_data = response.json()
                candidate = response_data["candidates"][0]
                text_content = candidate["content"]["parts"][0]["text"]
                return parse_json_response(text_content)
            else:
                last_error = f"Model {model} failed (Status {response.status_code}): {response.text}"
        except Exception as e:
            last_error = f"Model {model} exception: {str(e)}"
            
    raise Exception(f"All Gemini models failed. Details: {last_error}")


def analyze_with_llm(resume_text: str) -> dict:
    """Orchestrate LLM analysis trying OpenRouter first, falling back to Gemini."""
    schema = prepare_schema_for_gemini(ResumeProfile)
    
    prompt_text = f"""
You are an expert career decision analyst. Your job is to build the most accurate professional profile possible from the resume.

Instead of just asking "What is written in this resume?", evaluate "What does this resume tell us about this student's practical abilities, skills, and background?"

Analyze the resume text provided below and output a single JSON object matching the requested schema. Do not include any markdown styling like ```json or trailing text.

Here is the resume text:
---
{resume_text}
---

Structure your output to match the following JSON schema:
{json.dumps(schema, indent=2)}

Follow these guidelines for the extraction layers:
1. Personal Information: Extract contact details (name, email, phone, location, LinkedIn, GitHub, portfolio). Map their CURRENT education program (current year, college, degree, CGPA, graduation year) under current_education. If the current year is not explicitly written (e.g. '3rd year'), infer it based on their graduation year (e.g., if graduating in 2027 and current date is 2026, they are likely in their 3rd Year or Junior year).
2. Skills: For each skill, categorize it (e.g., Frontend, Backend, AI/ML, Database, DevOps, etc.), list evidence (such as projects or work experience names where they applied it), and estimate a confidence score (0.0 to 1.0) based on how strong the evidence is.
3. Projects: Extract detailed project records. Categorize projects into domain(s) (e.g., Backend, Frontend, AI, ML, IoT, Cloud, Blockchain, Cybersecurity, Data Science). Estimate complexity (High, Medium, Low) with a brief justification. Infer impact and key features.
4. Work Experience: Extract responsibilities, technologies, achievements, leadership initiatives, and business impact.
5. Education: Extract college/school records, degree, department, GPA/CGPA, coursework, and honors.
6. Certifications: Extract name, issuer, issue date, URL, and skills validated.
7. Extras: Pull out hackathons, open source contributions, competitions, volunteer work, leadership roles, positions of responsibility, achievements, awards, languages, interests, clubs, and society participations.

Analyze deep patterns rather than just doing text-matching. Provide honest and objective evaluations.
"""
    
    # Try OpenRouter first
    or_key = os.environ.get("OPENROUTER_API_KEY") or os.environ.get("OPENROUTER_APIKEY")
    if or_key:
        try:
            return call_openrouter_api(prompt_text)
        except Exception as exc:
            print(f"OpenRouter extraction failed: {exc}. Falling back to direct Gemini API...")
            
    # Try Gemini direct API next
    gemini_key = os.environ.get("GEMINI_API_KEY")
    if gemini_key:
        try:
            return call_gemini_api(prompt_text, schema)
        except Exception as exc:
            raise Exception(f"Both OpenRouter and direct Gemini APIs failed. Direct Gemini error: {str(exc)}")
            
    raise ValueError(
        "No API keys are configured. Please define OPENROUTER_API_KEY or GEMINI_API_KEY "
        "in your backend environment or a .env file."
    )


# ==========================================
# FILE EXTRACTION UTILITIES
# ==========================================

def extract_pdf(data: bytes) -> str:
    return "\n".join((page.extract_text() or "") for page in PdfReader(io.BytesIO(data)).pages)

def extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs]
    tables = [" | ".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows]
    return "\n".join(paragraphs + tables)


# ==========================================
# ROUTERS AND ENDPOINTS
# ==========================================

@app.get("/")
def index():
    if not (FRONTEND / "index.html").exists():
        return {"message": "Clarifuse Backend Running. Please run the frontend dev server via Vite during development."}
    return FileResponse(FRONTEND / "index.html")

@app.post("/api/resumes/extract")
async def extract_resume(file: UploadFile = File(...)):
    # Reload environment variables from disk dynamically
    load_dotenv()
    
    print("=" * 60)
    print("OPENROUTER =", repr(os.getenv("OPENROUTER_API_KEY")))
    print("GEMINI     =", repr(os.getenv("GEMINI_API_KEY")))
    print("=" * 60)
    
    # Verify at least one API key is present
    has_or = bool(os.environ.get("OPENROUTER_API_KEY") or os.environ.get("OPENROUTER_APIKEY"))
    has_gemini = bool(os.environ.get("GEMINI_API_KEY"))
    if not has_or and not has_gemini:
        paths_tried = [
            str((Path(__file__).parent / ".env").resolve()),
            str((Path(__file__).parent.parent / ".env").resolve()),
            str((Path.cwd() / ".env").resolve())
        ]
        debug_info = {
            "cwd": os.getcwd(),
            "paths_checked": {p: os.path.exists(p) for p in paths_tried},
            "env_keys": [k for k in os.environ.keys() if "KEY" in k or "API" in k or "OPEN" in k]
        }
        raise HTTPException(
            status_code=400,
            detail=(
                f"No API keys configured. Debug Info: {json.dumps(debug_info)}"
            )
        )

    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in {".pdf", ".doc", ".docx"}:
        raise HTTPException(415, "Please upload a PDF or DOCX resume.")
    if suffix == ".doc":
        raise HTTPException(415, "Legacy .doc files are not supported yet. Please save it as DOCX or PDF.")
    
    data = await file.read()
    if not data or len(data) > 10 * 1024 * 1024:
        raise HTTPException(413, "Use a non-empty resume smaller than 10 MB.")
    
    try:
        text = extract_pdf(data) if suffix == ".pdf" else extract_docx(data)
    except Exception as exc:
        raise HTTPException(422, "We couldn't read this file. Try exporting it as a text-based PDF or DOCX.") from exc
    
    if len(text.strip()) < 30:
        raise HTTPException(422, "We couldn't find readable text in this resume.")
    
    try:
        extraction = analyze_with_llm(text)
    except Exception as exc:
        raise HTTPException(500, f"Extraction failed: {str(exc)}")
    
    # Calculate word count for reference
    word_count = len(re.sub(r"\s+", " ", text).strip().split())
    
    return {
        "file_name": file.filename,
        "word_count": word_count,
        "extraction": extraction
    }
